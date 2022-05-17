from natasha import (
    Segmenter,
    MorphVocab,
    PER, LOC,
    NamesExtractor,
    NewsNERTagger,
    NewsEmbedding,
    NewsMorphTagger,
    Doc,
    DatesExtractor, AddrExtractor)
import hashlib
import sqlite3
from docx import Document
import re

segmenter = Segmenter()
morph_vocab = MorphVocab()
emb = NewsEmbedding()
morph_tagger = NewsMorphTagger(emb)
ner_tagger = NewsNERTagger(emb)

name_ex = NamesExtractor(morph_vocab)
ad_ex = AddrExtractor(morph_vocab)
date_ex = DatesExtractor(morph_vocab)

conn = sqlite3.connect('withOutFace.db', check_same_thread=False)
cur = conn.cursor()


def encrypt(line, change: int):
    texts = hashlib.md5(line.encode()).hexdigest()
    # 'абвгдеёжзийклмнопрстуфхцчшщъыьэюя'
    # 'abcdefghijklmnopqrstuvwxyz'
    string = 'abcdefghijklmnopqrstuvwxyz'
    line2 = list(string)
    line3 = list(string.upper())
    result_list = list(texts)
    for i in range(0, len(result_list)):
        if result_list[i] in line2 or result_list[i] in line3:
            if i < len(string) - 1 - change:
                result_list[i] = line2[i + change]
            elif i >= len(string) - 1 - change:
                result_list[i] = line2[i - len(string) - 1 + change]
    final = ''.join(result_list)
    return final


def markup(text, filename):
    replaceDic = []
    doc = Doc(text)
    doc.segment(segmenter)
    doc.tag_morph(morph_tagger)
    doc.tag_ner(ner_tagger)

    name = name_ex(text)
    i_name = 0
    for i in name:
        if i.fact.first is not None and i.fact.last is not None and i.fact.middle is not None:
            line = doc.text[i.start:i.stop]
            result = encrypt(line, 3)
            replaceDic.append([line, result, 'PER' + str(i_name)])
            i_name += 1
    date = date_ex(text)
    i_date = 0
    for j in date:
        line = doc.text[j.start:j.stop]
        result = encrypt(line, 3)
        replaceDic.append([line, result, 'DATA' + str(i_date)])
        i_date += 1
    i_addr = 0
    addr = ad_ex(text)
    for q in addr:
        line = doc.text[q.start:q.stop]
        result = encrypt(line, 3)
        replaceDic.append([line, result, 'LOC' + str(i_addr)])
        i_addr += 1


    i_other = 0
    sums = [r"[А-Я].{3,10} \w\.\w\.", r"\w\.\w\. \w{3,10}", r"\+\d{11}",
            r"8\d{10}", r"8-\d{3}-\d{3}-\d{2}-\d{2}", r"\+7\(\d{3}\)\d{7}",
            r"\+7 \(\d{3}\) \d{3}-\d{2}-\d{2}", r'(\b[\w.]+@+[\w.]+.+[\w.]\b)',
            r'(\b\+?[7,8](\s*\d{3}\s*\d{3}\s*\d{2}\s*\d{2})\b)', r'(\d{4}[\s\-]*\d{4}[\s\-]*\d{4}[\s\-]*\d{4})',
            r'(\b\w{5}\b\s\d{4}\b)', r'(\b\w{5}\b\s\d{6}\b)', r'\d{4}\D\d{6}']
    for qwe in sums:
        try:
            line = re.findall(qwe, text)[0]
            result = encrypt(line, 3)
            replaceDic.append([line, result, 'OTH' + str(i_other)])
            i_other += 1
        except:
            i_other += 1




    for i in replaceDic:
        text = text.replace(i[0], i[2])

    # a = re.search(r'\w\.\w\. [\n]\w{,15}', text)

    doc = Document('upload/fin.docx')
    paragraphs = doc.paragraphs
    for i in paragraphs:
        i.text = None
    doc.add_paragraph(text)
    doc.save('upload/' + filename)
    # with open('upload/test/keys.txt', 'w') as key:
    #     keys = ''
    #     for j in replaceDic:
    #         keys = keys + j[0] + '|' + j[1] + '|' + j[2] + '\n'
    #     key.write(keys)
    return text, replaceDic
